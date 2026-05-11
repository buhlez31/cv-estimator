"""PDF / DOCX → raw text + language detection."""

import io
import logging
import re
from pathlib import Path

import pdfplumber
import pypdf
from docx import Document

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the right parser by file extension."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_bytes)
    if suffix in (".docx", ".doc"):
        return _extract_docx(file_bytes)
    if suffix == ".txt":
        return _normalize(file_bytes.decode("utf-8", errors="replace"))
    raise ValueError(f"Unsupported file extension: {suffix!r}. Use PDF, DOCX, or TXT.")


def _extract_pdf(file_bytes: bytes) -> str:
    """Try pypdf first (pure-Python, stable). Fall back to pdfplumber if it
    yields no text — pypdf can fail silently on PDFs whose content streams
    use ToUnicode CMaps it doesn't decode, where pdfplumber/pdfminer succeed.

    pdfplumber is kept as a fallback rather than primary because its
    `pdfminer.six` backend has a history of native crashes on macOS for
    PDFs with embedded font tables.
    """
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages = [(page.extract_text() or "") for page in reader.pages]
        text = _normalize("\n".join(pages))
        if text.strip():
            return text
        logger.info("pypdf returned empty text; falling back to pdfplumber")
    except Exception as e:  # noqa: BLE001 — any pypdf failure → try fallback
        logger.warning("pypdf failed: %s; falling back to pdfplumber", e)

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return _normalize("\n".join(pages))


def _extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return _normalize("\n".join(paragraphs))


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


# --- Language detection (cs / en heuristic) ---
_CZ_CHARS = set("áčďéěíňóřšťúůýžÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ")
_CZ_STOPWORDS = {
    "a",
    "v",
    "se",
    "na",
    "je",
    "to",
    "že",
    "s",
    "z",
    "o",
    "do",
    "od",
    "pro",
    "nebo",
    "ale",
    "jako",
    "při",
    "byl",
    "byla",
    "bylo",
}
_EN_STOPWORDS = {
    "the",
    "of",
    "and",
    "to",
    "in",
    "is",
    "with",
    "for",
    "on",
    "as",
    "by",
    "an",
    "at",
    "from",
    "or",
}


def detect_language(text: str) -> str:
    """Return 'cs' or 'en'. Simple heuristic: diacritics + stopwords."""
    lower = text.lower()
    cz_char_hits = sum(1 for c in text if c in _CZ_CHARS)
    if cz_char_hits >= 20:
        return "cs"
    tokens = re.findall(r"[a-zá-ž]+", lower)
    if not tokens:
        return "en"
    cz_hits = sum(1 for t in tokens if t in _CZ_STOPWORDS)
    en_hits = sum(1 for t in tokens if t in _EN_STOPWORDS)
    return "cs" if cz_hits > en_hits else "en"
